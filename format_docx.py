from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

INPUT  = r"C:\Users\User\Downloads\medpaper\Case_Report_Epidermoid_Cyst_raw.docx"
OUTPUT = r"C:\Users\User\Downloads\medpaper\Case_Report_Epidermoid_Cyst_submission.docx"

doc = Document(INPUT)

# ── 1. Margins: 2.54 cm all sides ──────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── 2. Page numbers (bottom centre) ────────────────────────────────────────
def add_page_number(section):
    footer = section.footer
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.clear()
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

for section in doc.sections:
    add_page_number(section)

# ── 3. Default paragraph formatting ────────────────────────────────────────
HEADING_STYLES = {'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                  'Title', 'Subtitle'}

for para in doc.paragraphs:
    # Left alignment for body text
    if para.style.name not in HEADING_STYLES:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Double spacing + no space before/after
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)

    # Times New Roman 12 pt for all runs
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        # Preserve bold/italic from headings
        if para.style.name in HEADING_STYLES:
            run.font.bold = True

# ── 4. Tables (if any) ─────────────────────────────────────────────────────
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
