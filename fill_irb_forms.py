#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
fill_irb_forms.py  -  IRB forms filler -> medpaper/DOC/
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
from docx import Document
from docx.shared import Pt, RGBColor
import copy, os

SRC = r'C:\Users\User\Downloads'
DST = r'C:\Users\User\Downloads\medpaper\DOC'
os.makedirs(DST, exist_ok=True)

# ── 計畫資料 ──────────────────────────────────────────────────
TITLE_ZH   = '骨盆底巨大硬皮囊腫模擬直腸黏膜下惡性腫瘤：病例報告'
TITLE_EN   = 'A Large Epidermoid Cyst of the Pelvic Floor Mimicking Rectal Submucosal Malignancy: A Case Report'
PI_ZH      = '陳樞鴻'
PI_EN      = 'Chen Shu-Hung'
PI_INST    = '汐止國泰綜合醫院 直腸外科'
PI_TITLE   = '主治醫師'
PI_EMAIL   = 'cgh07668@cgh.org.tw'
COPI_ZH    = '陳紹寬'
COPI_DEPT  = '泌尿科'
COPI_TITLE = '主治醫師'
RA1_ZH     = '王子建'
RA2_ZH     = '吳美智'
RA_DEPT    = '直腸外科'
RA_TITLE   = '專科護理師'
DATE_END   = '民國116年12月31日'
SUBJ_NUM   = '1'
COI_TRUST  = '否'

ABSTRACT_ZH = (
    '骨盆底硬皮囊腫（epidermoid cyst）為極罕見良性病灶（文獻報告少於20例），'
    '電腦斷層掃描（CT）上易與惡性腫瘤混淆，造成診斷困難。一名61歲男性因下腹疼痛、'
    '血尿、排便習慣改變及裡急後重就醫。對比增強CT顯示左側低位直腸旁10公分腫塊，'
    '伴見疑似轉移性淋巴結，初步鑑別診斷包括胃腸道基質瘤（GIST）、淋巴瘤及神經內分泌腫瘤。'
    '三項術前發現協助導向正確診斷：大腸鏡無黏膜病灶、腫瘤標記（CEA、AFP、CA19-9）均在正常範圍，'
    '以及CT腫塊未見對比增強。泌尿科會診CT判讀及術中確認左側第三級水腎（SFU分級），'
    '術前預防性置入雙J管。患者接受剖腹手術完整切除骨盆底腫塊，病理報告確認為硬皮囊腫，'
    '術後恢復良好，第9天出院。骨盆底硬皮囊腫應列入骨盆腔腫塊之鑑別診斷，'
    '完整手術切除具治癒性，預後良好。'
)
ABSTRACT_EN = (
    'Epidermoid cysts of the pelvic floor are rare benign lesions that can closely '
    'mimic malignancy on CT, posing significant diagnostic challenges. A 61-year-old '
    'male presented with lower abdominal pain, hematuria, altered bowel habits, and '
    'rectal tenesmus. Contrast-enhanced CT revealed a 10 cm pelvic mass with apparent '
    'lymphadenopathy, raising suspicion for GIST, lymphoma, or neuroendocrine tumor. '
    'Three pre-operative findings redirected the diagnosis: negative colonoscopy, normal '
    'tumor markers (CEA, AFP, CA19-9), and absent mass enhancement on CT. Urology '
    'consultation identified Grade 3 left hydronephrosis (SFU grading) on CT review, '
    'prompting prophylactic double-J ureteral stenting. Laparotomy with complete excision '
    'confirmed an epidermoid cyst on histopathology. The patient was discharged on '
    'postoperative day 9 without complications. When a non-enhancing extrinsic pelvic '
    'mass is accompanied by negative colonoscopy and normal tumor markers, epidermoid '
    'cyst should be included in the differential. Complete surgical excision is curative.'
)
KEYWORDS_ZH = '硬皮囊腫；骨盆底；骨盆腫塊；鑑別診斷；剖腹手術'
KEYWORDS_EN = 'epidermoid cyst; pelvic floor; pelvic mass; differential diagnosis; laparotomy'
OBJECTIVE   = '報告一例骨盆底巨大硬皮囊腫在CT上模擬直腸黏膜下惡性腫瘤之罕見案例，描述術前診斷重新定向之關鍵特徵及手術處置策略'
STUDY_DESIGN = '單一機構回溯性病例報告（Retrospective single-case report）'

# ── 工具函式 ──────────────────────────────────────────────────
def _set_para_text(para, text):
    """Replace paragraph text while keeping first run's formatting."""
    if para.runs:
        para.runs[0].text = text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(text)

def set_cell(cell, text):
    """Set first paragraph of cell to text."""
    _set_para_text(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        for r in p.runs:
            r.text = ''

def replace_runs(para, old, new):
    """Replace old→new across all runs in a paragraph."""
    full = para.text
    if old not in full:
        return False
    rebuilt = full.replace(old, new)
    for r in para.runs:
        r.text = ''
    if para.runs:
        para.runs[0].text = rebuilt
    else:
        para.add_run(rebuilt)
    return True

def replace_in_doc(doc, old, new):
    """Replace text everywhere in document."""
    count = 0
    for p in doc.paragraphs:
        if replace_runs(p, old, new):
            count += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if replace_runs(p, old, new):
                        count += 1
    return count

def uniq_cells(row):
    """Return deduplicated cells from a row (handle merged cells)."""
    seen, result = set(), []
    for c in row.cells:
        cid = id(c._tc)
        if cid not in seen:
            seen.add(cid)
            result.append(c)
    return result

# ════════════════════════════════════════════════════════════
# 表格 10-4：個案報告審查申請表
# ════════════════════════════════════════════════════════════
def fill_10_4():
    src = os.path.join(SRC, '65_10-4.docx')
    dst = os.path.join(DST, '10-4_審查申請表_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]
    t1 = doc.tables[1]

    # ─ Table 0 ─
    # R1 中文計畫名稱
    set_cell(uniq_cells(t0.rows[1])[1], TITLE_ZH)
    # R2 英文計畫名稱
    set_cell(uniq_cells(t0.rows[2])[1], TITLE_EN)
    # R3 主持人姓名（3 uniq cells: label | 中文名 | 英文名）
    set_cell(uniq_cells(t0.rows[3])[1], f'中文姓名：{PI_ZH}')
    set_cell(uniq_cells(t0.rows[3])[-1], f'英文姓名：{PI_EN}')
    # R4 服務機構/職稱
    set_cell(uniq_cells(t0.rows[4])[1], f'服務機構/單位：{PI_INST}')
    set_cell(uniq_cells(t0.rows[4])[-1], f'職稱：{PI_TITLE}')
    # R5 電話/email
    set_cell(uniq_cells(t0.rows[5])[1], '聯絡電話：（請主持人填入院內分機）')
    set_cell(uniq_cells(t0.rows[5])[-1], f'E-mail：{PI_EMAIL}')
    # R6 執行期間
    set_cell(uniq_cells(t0.rows[6])[1],
             f'IRB許可書核准日 至 {DATE_END}')
    # R7 研究對象人數
    set_cell(uniq_cells(t0.rows[7])[1], f'(國內 {SUBJ_NUM} 人，本院 {SUBJ_NUM} 人)')
    # R8 經費贊助
    set_cell(uniq_cells(t0.rows[8])[1], '■　無   □　有，名稱：')
    # R9 研究成員 header — 已保留
    # R10 共/協同主持人
    set_cell(uniq_cells(t0.rows[10])[1],
             f'{COPI_ZH}　　{COPI_DEPT}　　（院內分機）　　（email）')
    # R11 研究助理/護士
    set_cell(uniq_cells(t0.rows[11])[1],
             f'{RA1_ZH}  直腸外科  （分機）  （email）／{RA2_ZH}  直腸外科  （分機）  （email）')
    # R12 聯絡人
    cells12 = uniq_cells(t0.rows[12])
    if len(cells12) > 1:
        set_cell(cells12[1], f'{RA1_ZH}  直腸外科  （分機）  （email）')
    # R13-R15 文件版本
    set_cell(uniq_cells(t0.rows[13])[1], '中文摘要（計畫書）')
    set_cell(uniq_cells(t0.rows[14])[1], 'v1.0')
    set_cell(uniq_cells(t0.rows[15])[1], '2026年06月02日')

    # ─ Table 1 ─
    # R0 是否在其他單位送審
    for p in t1.rows[0].cells[1].paragraphs:
        replace_runs(p, '□  無', '■  無')
    # R1 研究設計
    for p in t1.rows[1].cells[1].paragraphs:
        replace_runs(p, '□ 病例報告', '■ 病例報告')
    # R2 研究對象
    cells_r2 = uniq_cells(t1.rows[2])
    for c in cells_r2:
        for p in c.paragraphs:
            replace_runs(p, '□ 病患', '■ 病患')
    # R3 年齡
    for p in t1.rows[3].cells[1].paragraphs:
        replace_runs(p, '________歲～________歲', '61歲～61歲')
    # R4 特殊族群
    for p in t1.rows[4].cells[1].paragraphs:
        replace_runs(p, '(弱勢團體須排除)', '成年男性，無特殊弱勢族群（不符合任何弱勢定義）')
    # R5 資料保密 — 勾電腦資訊 + 以編號識別
    for p in t1.rows[5].cells[1].paragraphs:
        replace_runs(p, '□電腦資訊', '■電腦資訊')
        replace_runs(p, '□以編號識別(請說明)：', '■以編號識別（病歷號已去識別化）')
        replace_runs(p, '□所有資料上鎖：', '■所有資料上鎖：')
        replace_runs(p, '□地點(請說明', '■地點（直腸外科辦公室）')

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-5：中英文摘要表
# ════════════════════════════════════════════════════════════
def fill_10_5():
    src = os.path.join(SRC, '65_10-5.docx')
    dst = os.path.join(DST, '10-5_中英文摘要表_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]
    rows = t0.rows

    # R0 中文計畫名稱
    set_cell(uniq_cells(rows[0])[0], f'計畫名稱：(中文)\n{TITLE_ZH}')
    # R1 英文計畫名稱
    set_cell(uniq_cells(rows[1])[0], f'計畫名稱：(英文)\n{TITLE_EN}')
    # R2 試驗目的
    set_cell(uniq_cells(rows[2])[0], f'試驗目的：\n{OBJECTIVE}')
    # R3 研究設計
    set_cell(uniq_cells(rows[3])[0], f'研究設計：\n{STUDY_DESIGN}')
    # R4 中文摘要（兩行 R4+R5 for label+content）
    set_cell(uniq_cells(rows[4])[0],
             f'試驗計劃摘要（中文）：\n{ABSTRACT_ZH}')
    # R6 中文關鍵字
    set_cell(uniq_cells(rows[6])[0], f'關鍵字：{KEYWORDS_ZH}')
    # R7 英文摘要
    set_cell(uniq_cells(rows[7])[0],
             f'試驗計劃摘要（英文）：\n{ABSTRACT_EN}')
    # R9 英文關鍵字
    set_cell(uniq_cells(rows[9])[0], f'Key Word：{KEYWORDS_EN}')

    # 頁腳備註：版本日期
    doc.sections[0].footer.paragraphs[0].text if doc.sections[0].footer.paragraphs else None
    # 在最後一段加版本
    last_para = doc.paragraphs[-1]
    if 'v1.0' not in last_para.text:
        last_para.add_run('　　　　版本：v1.0　日期：2026-06-02')

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-6：病人資料提供同意書
# ════════════════════════════════════════════════════════════
def fill_10_6():
    src = os.path.join(SRC, '65_10-6.docx')
    dst = os.path.join(DST, '10-6_病人同意書_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]

    # R0：填入計畫名稱、執行單位、主持人
    r0_text = (
        '本試驗已通過國泰綜合醫院人體試驗審查委員會審查，計畫編號：CGH-（待IRB給號）\n'
        f'計劃名稱：{TITLE_ZH}\n'
        f'計劃執行單位：{PI_INST}\n'
        f'計劃主持人姓名：{PI_ZH}　　職稱：{PI_TITLE}　　聯絡電話：（院內分機）\n'
        '通信地址：新北市汐止區建成路59巷2號\n'
        '24小時緊急聯絡電話：（主持人行動電話，請主持人自行填入）'
    )
    _set_para_text(t0.rows[0].cells[0].paragraphs[0], r0_text)

    # R1：受試者資料（個人資料由臨床團隊填入）
    r1_text = (
        '受試者姓名：________________（由主持人從病歷填入）\n'
        '性別：男　　出生日期：________________　　病歷號碼：________________\n'
        '通信地址：________________\n'
        '聯絡電話：________________'
    )
    _set_para_text(t0.rows[1].cells[0].paragraphs[0], r1_text)

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-7-1：財務利益評估說明表（主持人代表填一份）
# ════════════════════════════════════════════════════════════
def fill_10_7_1():
    src = os.path.join(SRC, '65_10-7-1_converted.docx')
    dst = os.path.join(DST, '10-7-1_財務評估說明表_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]

    rows = t0.rows
    # R0 IRB編號（待填）/ 計畫編號（待填）
    set_cell(uniq_cells(rows[0])[0], 'IRB編號：（待 IRB 秘書處給號）')
    set_cell(uniq_cells(rows[0])[-1], '計畫編號：CGH-（待給號）')
    # R1 計畫主持人 / 單位
    set_cell(uniq_cells(rows[1])[0], f'計畫主持人：{PI_ZH}')
    set_cell(uniq_cells(rows[1])[-1], f'單位：{PI_INST}')
    # R2 聯絡人 / 聯絡方式
    set_cell(uniq_cells(rows[2])[0], f'聯絡人：{RA1_ZH}')
    set_cell(uniq_cells(rows[2])[-1], '聯絡方式：（分機）（email）')
    # R3 計畫名稱
    set_cell(uniq_cells(rows[3])[0], f'計畫名稱：{TITLE_ZH}')
    # R4 試驗委託者
    set_cell(uniq_cells(rows[4])[0], '試驗委託者：無（本院自主研究）')

    # 問題與回答（A欄 Q1–Q6 answer rows at odd positions 7,9,11,13,15,17）
    answers = [
        # Q1 學術價值
        ('請說明：',
         '本病例為極罕見骨盆底硬皮囊腫（文獻報告少於20例），具有高度CT惡性腫瘤假象（疑似轉移性淋巴結）。'
         '此影像特徵及術前三聯徵（大腸鏡無黏膜病灶、腫瘤標記正常、腫塊無強化）對臨床醫師辨識此類罕見病灶具有重要教學價值，'
         '可降低誤診率，具有一定學術貢獻。'),
        # Q2 受試者風險
        ('請說明：',
         '本研究為回溯性病例報告，不涉及任何額外處置或介入，僅整理已完成之診療病歷資料進行學術發表。'
         '受試者風險極低，唯個人隱私保護需妥善處理（病歷號及影像均已去識別化）。'),
        # Q3 財務利益種類
        ('請說明：',
         '無。本研究無任何廠商贊助、委託費用、股權或智慧財產權相關財務利益。所有研究人員均聲明無顯著財務利益。'),
        # Q4 是否影響研究
        ('請說明：',
         '不適用。本研究無任何財務利益/非財務關係，不存在利益衝突之疑慮，無法影響研究執行或結果。'),
        # Q5 獨特能力
        ('請說明：',
         '本案例由汐止國泰綜合醫院直腸外科與泌尿科合作診治，研究團隊為直接參與診斷與手術之醫護人員，'
         '為執行本病例報告之最適當人選。'),
        # Q6 主管職權關係
        ('請說明：',
         '不適用。持有財務利益者與研究人員間無涉及財務利益之上下屬督導關係。'),
    ]
    answer_row_indices = [7, 9, 11, 13, 15, 17]
    for idx, (label, ans) in zip(answer_row_indices, answers):
        if idx < len(rows):
            cells = uniq_cells(rows[idx])
            if len(cells) >= 2:
                set_cell(cells[1], ans)
            elif cells:
                set_cell(cells[0], f'{label}\n{ans}')

    # B欄 處置計畫 (Table 1)
    t1 = doc.tables[1]
    for p in t1.rows[1].cells[0].paragraphs:
        replace_runs(p, '□撤除所有的顯著財務利益/非財務關係。', '□撤除所有的顯著財務利益/非財務關係。')
    # 無需採取處置 — 直接聲明
    _set_para_text(t1.rows[1].cells[0].paragraphs[0],
                   '■ 無任何顯著財務利益/非財務關係，不需採取處置措施。\n'
                   '本研究為無贊助之回溯性病例報告，所有研究人員均無需申報之財務利益。')

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-7-2：財務利益申報表（每人一份，共 4 份）
# ════════════════════════════════════════════════════════════
def fill_10_7_2_one(person_name, person_dept, person_title, is_pi=False, suffix=''):
    src = os.path.join(SRC, '65_10-7-2_converted.docx')
    safe = person_name.replace('/', '_')
    dst = os.path.join(DST, f'10-7-2_財務申報表_{safe}_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]
    rows = t0.rows

    # R0 申報類型 — 勾新研究計畫申請
    for p in rows[0].cells[0].paragraphs:
        replace_runs(p, '□新研究計畫申請', '■新研究計畫申請')

    # R1 計畫名稱
    set_cell(uniq_cells(rows[1])[0], f'計畫名稱：{TITLE_ZH}')
    if len(uniq_cells(rows[1])) > 1:
        set_cell(uniq_cells(rows[1])[1], f'計畫名稱：{TITLE_ZH}')

    # R2 計畫主持人
    set_cell(uniq_cells(rows[2])[0], f'計畫主持人：{PI_ZH}')
    set_cell(uniq_cells(rows[2])[-1], '連絡電話：（院內分機）')

    # R3 試驗委託者
    set_cell(uniq_cells(rows[3])[0], '試驗委託者：無')
    set_cell(uniq_cells(rows[3])[-1], '連絡電話：不適用')

    # R4 潛在之試驗機構財務利益 — 勾否
    for p in rows[4].cells[0].paragraphs:
        replace_runs(p, '□ 否', '■ 否')

    # R5 A欄標題 — 保留
    # R6 A欄聲明內文 — 本人資料
    person_decl = (
        f'本人茲聲明：\n'
        f'申報人：{person_name}　單位：{person_dept}　職稱：{person_title}\n'
        f'本人、本人配偶與未成年子女，目前無任何依國泰綜合醫院政策必須申報之'
        f'「顯著財務利益」及「非財務關係」；\n'
        f'本研究為無經費贊助之回溯性病例報告，研究人員無任何財務利益需申報。\n'
        f'申報人簽名：________________　日期：　　年　　月　　日'
    )
    _set_para_text(rows[6].cells[0].paragraphs[0], person_decl)

    # PI額外聲明 (R11/R12)
    if is_pi and len(rows) > 11:
        pi_stmt = (
            f'計畫主持人之聲明：\n'
            f'顯著財務利益/非財務關係之所有研究相關人員，已詳列並提出本表。\n'
            f'所有研究人員（{PI_ZH}、{COPI_ZH}、{RA1_ZH}、{RA2_ZH}）均填妥A欄聲明無顯著財務利益。\n'
            f'計畫主持人簽名：________________　日期：　　年　　月　　日'
        )
        _set_para_text(rows[11].cells[0].paragraphs[0], pi_stmt)

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-1：收件表格
# ════════════════════════════════════════════════════════════
def fill_10_1():
    src = os.path.join(SRC, '65_10-1_converted.docx')
    dst = os.path.join(DST, '10-1_收件表格_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]
    rows = t0.rows

    # R1 計畫名稱中文
    cells = uniq_cells(rows[1])
    if len(cells) > 1:
        set_cell(cells[1], TITLE_ZH)

    # R2 計畫名稱英文
    cells = uniq_cells(rows[2])
    if len(cells) > 1:
        set_cell(cells[1], TITLE_EN)

    # R3 計劃主持人
    cells = uniq_cells(rows[3])
    if len(cells) > 1:
        set_cell(cells[1],
                 f'姓名：{PI_ZH}　科別：直腸外科　職稱：{PI_TITLE}　聯絡電話：（分機）')

    # R4 協同主持人
    cells = uniq_cells(rows[4])
    if len(cells) > 1:
        set_cell(cells[1],
                 f'姓名：{COPI_ZH}　科別：{COPI_DEPT}　職稱：{COPI_TITLE}　聯絡電話：（分機）')

    # R5 研究人員
    cells = uniq_cells(rows[5])
    if len(cells) > 1:
        set_cell(cells[1],
                 f'姓名：{RA1_ZH}　科別：{RA_DEPT}　職稱：{RA_TITLE}　聯絡電話：（分機）\n'
                 f'姓名：{RA2_ZH}　科別：{RA_DEPT}　職稱：{RA_TITLE}　聯絡電話：（分機）')

    # R6 聯絡人
    cells = uniq_cells(rows[6])
    if len(cells) > 1:
        set_cell(cells[1],
                 f'姓名：{RA1_ZH}　聯絡電話：（分機）　e-mail：（請填入）')

    # R8–R16 文件備妥欄位 (V) — 在 C3 備妥欄打 V
    doc_rows = [8, 9, 10, 11, 12, 13, 14, 15, 16]
    for ri in doc_rows:
        if ri < len(rows):
            cells = uniq_cells(rows[ri])
            # 找備妥欄（通常是第3或第4個唯一cell）
            for ci, c in enumerate(cells):
                if 'V' in c.text or '備妥' in c.text or ci == 3:
                    set_cell(c, 'V')
                    break

    doc.save(dst)
    print(f'✓ {dst}')

# ════════════════════════════════════════════════════════════
# 表格 10-2：履歷表（每人一份，共 4 份）
# ════════════════════════════════════════════════════════════
def fill_10_2_one(person_name, gender, dept, inst, title, role):
    src = os.path.join(SRC, '65_10-2_converted.docx')
    safe = person_name.replace('/', '_')
    dst = os.path.join(DST, f'10-2_履歷_{safe}_填.docx')
    doc = Document(src)
    t0 = doc.tables[0]
    rows = t0.rows

    # R0 類別
    role_map = {'PI': '( ✓ ) 主持人', 'CoPI': '( ✓ ) 協同主持人', 'RA': '( ✓ ) 研究人員'}
    set_cell(uniq_cells(rows[0])[1], role_map.get(role, role))

    # R1 姓名/性別/出生
    cells1 = uniq_cells(rows[1])
    if len(cells1) >= 3:
        set_cell(cells1[0], f'姓名：{person_name}')
        set_cell(cells1[1], f'性別：{gender}')
        set_cell(cells1[2], '出生年月日：（請自行填入）')

    # R2 學歷標題保留
    # R9 現任
    cells9 = uniq_cells(rows[9])
    if cells9:
        set_cell(cells9[0], f'現任：{inst}　{dept}　{title}')

    # R12 近五年著作（空白待填）
    cells12 = uniq_cells(rows[12])
    if cells12:
        role_zh = {'PI': '通訊作者', 'CoPI': '共同作者', 'RA': '共同作者'}.get(role, '')
        set_cell(cells12[0],
                 f'近五年相關著作及研究報告名稱：（請自行填入，含本次投稿論文）\n'
                 f'【本次投稿】{TITLE_EN}，Journal of the Formosan Medical Association（Under Review）')

    doc.save(dst)
    print(f'✓ {dst}')


# ════════════════════════════════════════════════════════════
# 主程式
# ════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('開始填寫 IRB 表格...\n')

    fill_10_4()
    fill_10_5()
    fill_10_6()
    fill_10_7_1()

    # 10-7-2：四位研究人員各一份
    fill_10_7_2_one(PI_ZH,    '直腸外科', PI_TITLE,    is_pi=True)
    fill_10_7_2_one(COPI_ZH,  COPI_DEPT,  COPI_TITLE,  is_pi=False)
    fill_10_7_2_one(RA1_ZH,   RA_DEPT,    RA_TITLE,    is_pi=False)
    fill_10_7_2_one(RA2_ZH,   RA_DEPT,    RA_TITLE,    is_pi=False)

    # 10-1：收件表格
    fill_10_1()

    # 10-2：四位研究人員履歷各一份
    fill_10_2_one(PI_ZH,   '男', '直腸外科', PI_INST, PI_TITLE,    'PI')
    fill_10_2_one(COPI_ZH, '男', COPI_DEPT,  PI_INST, COPI_TITLE,  'CoPI')
    fill_10_2_one(RA1_ZH,  '男', RA_DEPT,    PI_INST, RA_TITLE,    'RA')
    fill_10_2_one(RA2_ZH,  '女', RA_DEPT,    PI_INST, RA_TITLE,    'RA')

    print(f'\n✅ 完成！輸出目錄：{DST}')
    print('\n⚠️  以下欄位需人工補填：')
    print('  • 各人院內電話分機')
    print('  • 各人 email（非陳樞鴻者）')
    print('  • 10-6 病人姓名、出生日期、病歷號')
    print('  • 10-2 各人出生年月日、學歷、曾任經歷')
    print('  • 所有表格的親筆簽名欄位')
    print('  • IRB 案號（待秘書處給號後填入 10-1、10-4、10-7-1）')
